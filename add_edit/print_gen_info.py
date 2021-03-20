def print_info(cursor, file_number, folders):
    import docx
    from datetime import date
    import helper_function.pccm_names as pccm_names
    import os.path
    module_name = ["bio_info", "nut_supplements", "phys_act", "habits", "family_details", "med_history", "cancer_history",
                    "family_cancer", "det_by", "breast_symptoms"]
    folder = folders +'/Gen_Info_Docs'
    file_name = "Folder_" + (file_number.replace("/", "_")) + ".docx"
    path = os.path.join(folder, file_name)
    p_date_to_print = date.today().strftime('%d-%b-%Y')
    doc = docx.Document()
    doc.add_paragraph(("File Number " + file_number), style="Title")
    doc.add_paragraph(("Document Created on " + p_date_to_print), style="Quote")
    for index in range(0, len(module_name)):
        col_titles = pccm_names.names_info(module_name[index])
        columns = ", ".join(col_titles)
        sql = 'SELECT ' + columns + ' FROM Patient_Information_History WHERE File_number = \'' + file_number + "'"
        data = cursor.execute(sql)
        data_file = data.fetchall()
        data_print = list(data_file[0])
        text_titles = pccm_names.info_print_all(module_name[index])
        for i in range(0, len(data_print)):
            p = doc.add_paragraph((text_titles[i] + ":  "), style="List Bullet")
            (p.add_run(data_print[i])).bold = True
        doc.add_paragraph()
    doc.save(path)